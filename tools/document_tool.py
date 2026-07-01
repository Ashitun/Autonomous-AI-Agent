from docx import Document
from datetime import datetime
import re
import os


def generate_document(goal: str, results: list):

    document = Document()

    document.add_heading(goal, level=1)

    for item in results:

        document.add_heading(item["task"], level=2)

        document.add_paragraph(item["content"])

    os.makedirs("generated_docs", exist_ok=True)

    safe_goal = re.sub(r'[^a-zA-Z0-9 ]', '', goal)
    safe_goal = safe_goal.lower().replace(" ", "_")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    filename = f"{safe_goal}_{timestamp}.docx"

    file_path = os.path.join(
        "generated_docs",
        filename
    )

    document.save(file_path)

    return file_path